import os
import random
from docx import Document
import tkinter as tk
from tkinter import messagebox

class Genquestionqcm:
    def __init__(self, question, reponses, juste):
        self.question = question
        self.reponses = reponses  # liste des réponses
        self.juste = juste

    def __str__(self):
        return f"question: {self.question}\nreponses: {', '.join(self.reponses)}\nreponse correcte: {self.juste}\n"

def lectureFichier(nomDeFichier):
    """Permet de récupérer les questions d'un fichier .txt"""
    listeQuestions = []
    with open(nomDeFichier, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        index = 0
        while index < len(lines):
            line = lines[index].strip()
            if line.endswith("?"):
                question = line
                rep_juste = lines[index+5].strip()
                reponses = [lines[index+1].strip(), lines[index+2].strip(), lines[index+3].strip(), lines[index+4].strip()]
                listeQuestions.append(Genquestionqcm(question, reponses, rep_juste))
                index += 6
            else:
                index += 1
        return listeQuestions

def qcmdocx(liste):
    qcm = Document()
    nbquestion = len(liste)
    for x in liste:
        qcm.add_paragraph(x.question)
        qcm.add_paragraph(x.reponses[0])
        qcm.add_paragraph(x.reponses[1])
        qcm.add_paragraph(x.reponses[2])
        qcm.add_paragraph(x.reponses[3])
        qcm.add_paragraph()
    if nbquestion == 5:
        tableau = qcm.add_table(rows=2, cols=5)
        for i in range(2):  # Parcourir les lignes
            for j in range(5):  # Parcourir les colonnes
                cellule = tableau.cell(i, j)
                if i == 0:
                    cellule.text = f'{j+1}'
    elif nbquestion == 10:
        tableau = qcm.add_table(rows=2, cols=11)
        for i in range(2):  # Parcourir les lignes
            for j in range(11):  # Parcourir les colonnes
                cellule = tableau.cell(i, j)
                if i == 0 and j < 5:
                    cellule.text = f'{j+1}'
                if i == 0 and j > 5:
                    cellule.text = f'{j}'
    elif nbquestion == 15:
        tableau1 = qcm.add_table(rows=2, cols=11)
        for i in range(2):  # Parcourir les lignes
            for j in range(11):  # Parcourir les colonnes
                cellule = tableau1.cell(i, j)
                if i == 0 and j < 5:
                    cellule.text = f'{j+1}'
                if i == 0 and j > 5:
                    cellule.text = f'{j}'
        tableau2 = qcm.add_table(rows=2, cols=5)
        for i in range(2):  # Parcourir les lignes
            for j in range(5):  # Parcourir les colonnes
                cellule = tableau2.cell(i, j)
                if i == 0:
                    cellule.text = f'{j + 11}'
    else:
        tableau = qcm.add_table(rows=2, cols=11)
        for i in range(2):  # Parcourir les lignes
            for j in range(11):  # Parcourir les colonnes
                cellule = tableau.cell(i, j)
                if i == 0 and j < 5:
                    cellule.text = f'{j+1}'
                if i == 0 and j > 5:
                    cellule.text = f'{j}'
        tableau2 = qcm.add_table(rows=2, cols=11)
        for i in range(2):  # Parcourir les lignes
            for j in range(11):  # Parcourir les colonnes
                cellule = tableau2.cell(i, j)
                if i == 0 and j < 5:
                    cellule.text = f'{j+11}'
                if i == 0 and j > 5:
                    cellule.text = f'{j+10}'
    qcm.save(r'D:\COURS\genQCM\quiz_questions.docx')

def repdocx(reponses):
    rep = Document()
    rep.add_paragraph(reponses)
    rep.save(r'D:\COURS\genQCM\correction.docx')

def creation_sujets(liste, nb_suj, nb_qst):
    sujets = []
    reponses_sujets = []
    for i in range(0, nb_suj):
        sujet = []
        reponses = ""
        while len(sujet) < nb_qst:
            choix_question = random.randint(0, len(liste)-1)
            if liste[choix_question] not in sujet:
                sujet.append(liste[choix_question])
                reponses += str(liste[choix_question].juste)
        sujets.append(sujet)
        reponses_sujets.append(reponses)
    return (sujets, reponses_sujets)

def on_submit():
    try:
        nb_suj = int(entry_sujets.get())
        nb_qst = int(entry_questions.get())
        if nb_qst % 5 != 0:
            messagebox.showerror("Erreur", "Le nombre de questions doit être un multiple de 5")
            return
        sujets, reponses = creation_sujets(questions, nb_suj, nb_qst)
        qcmdocx(sujets[0])
        repdocx(reponses[0])
        messagebox.showinfo("Succès", "Les fichiers ont été générés avec succès!")
    except ValueError:
        messagebox.showerror("Erreur", "Veuillez entrer des nombres valides")

# Interface Tkinter
root = tk.Tk()
root.title("Générateur de QCM")

tk.Label(root, text="Nombre de sujets:").grid(row=0, column=0)
entry_sujets = tk.Entry(root)
entry_sujets.grid(row=0, column=1)

tk.Label(root, text="Nombre de questions (multiples de 5):").grid(row=1, column=0)
entry_questions = tk.Entry(root)
entry_questions.grid(row=1, column=1)

tk.Button(root, text="Générer", command=on_submit).grid(row=2, column=0, columnspan=2)

# Charger les questions
questions = lectureFichier(r"D:\COURS\genQCM\QCM_culture_generale.txt")

root.mainloop()